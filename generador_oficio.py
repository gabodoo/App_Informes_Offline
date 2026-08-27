import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
import sys

def get_base_dir() -> Path:
    if getattr(sys, 'frozen', False):
        return Path(sys.executable).parent
    return Path(__file__).resolve().parent

class GeneradorOficio:
    """Clase para la generación del Oficio a partir de su plantilla."""

    def __init__(self):
        self.ruta_plantilla = get_base_dir() / "plantillas" / "plantilla_oficio.docx"

    def generar(self, contexto: dict, log_callback=None) -> Path:
        def log(msg):
            if log_callback:
                log_callback(msg)

        if not self.ruta_plantilla.exists():
            log(f"No se encontró la plantilla del oficio en {self.ruta_plantilla}")
            raise FileNotFoundError(f"Falta plantilla: {self.ruta_plantilla}")

        log("Cargando plantilla de oficio...")
        doc = Document(self.ruta_plantilla)
        
        # 1. Reemplazos directos
        sigedo_corto = contexto.get("SIGEDO_CORTO", "")
        reemplazo_sigedo = str(contexto.get("NUMERO_SIGEDO", ""))
        n_bec = str(contexto.get("NOMBRES_BECARIO", "")).strip()
        a_bec = str(contexto.get("APELLIDOS_BECARIO", "")).strip()
        reemplazo_nombres = f"{n_bec} {a_bec}".strip()
        if not reemplazo_nombres:
            reemplazo_nombres = str(contexto.get("NOMBRES_Y_APELLIDOS_VALIDADOS", "")).upper()
        reemplazo_informe = f"Informe N° {contexto.get('NUMERO_INFORME_GENERAR', '')}-2026-MINEDU/VMGI-PRONABEC-DIBEC-SUS"
        reemplazo_fecha = f"Escrito de fecha {contexto.get('FECHA_SOLICITUD_TEXTO', '')}"
        trato = str(contexto.get("TRATO_GENERO", "Señorita"))
        
        patrones_reemplazo = [
            ("57897-2026", reemplazo_sigedo),
            ("Señorita", trato),
            ("MAINIA ROSITA SEKUTA KAYACH", reemplazo_nombres),
            ("Escrito de fecha 17 de julio de 2026", reemplazo_fecha),
            ("Informe N° 6541-2026-MINEDU/VMGI-PRONABEC-DIBEC-SUS", reemplazo_informe),
            ("Informe N° 6541-2026-MINEDU/VMGI-PRONABEC-DIBEC", reemplazo_informe),
            ("Informe N° 6541", reemplazo_informe),
            ("57897", sigedo_corto),
        ]

        def _reemplazar_texto_run(run):
            texto_nuevo = run.text
            for tag, val in patrones_reemplazo:
                if tag in texto_nuevo:
                    texto_nuevo = texto_nuevo.replace(tag, val)
            
            if texto_nuevo != run.text:
                run.text = texto_nuevo

        def procesar_parrafo(p):
            if not p.text:
                return
                
            texto_pre = p.text
            for run in p.runs:
                _reemplazar_texto_run(run)

            # Fallback a nivel de párrafo
            if p.text == texto_pre:
                texto_orig = p.text
                texto_nuevo = texto_orig
                for tag, val in patrones_reemplazo:
                    if tag in texto_nuevo:
                        texto_nuevo = texto_nuevo.replace(tag, val)
                
                if texto_nuevo != texto_orig:
                    fmt_guardado = {}
                    if p.runs:
                        r0 = p.runs[0]
                        fmt_guardado = {
                            "name": r0.font.name,
                            "size": r0.font.size,
                            "bold": r0.bold,
                            "italic": r0.italic,
                            "underline": r0.underline,
                        }
                    p.text = texto_nuevo
                    if p.runs and fmt_guardado:
                        r = p.runs[0]
                        if fmt_guardado.get("name"):
                            r.font.name = fmt_guardado["name"]
                        if fmt_guardado.get("size"):
                            r.font.size = fmt_guardado["size"]
                        if fmt_guardado.get("bold") is not None:
                            r.bold = fmt_guardado["bold"]
                        if fmt_guardado.get("italic") is not None:
                            r.italic = fmt_guardado["italic"]
                        if fmt_guardado.get("underline") is not None:
                            r.underline = fmt_guardado["underline"]

        # Procesar documento entero
        for p in doc.paragraphs:
            procesar_parrafo(p)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        procesar_parrafo(p)

        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                if header:
                    for p in header.paragraphs:
                        procesar_parrafo(p)
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    procesar_parrafo(p)

            for footer in [section.footer, section.first_page_footer]:
                if footer:
                    for p in footer.paragraphs:
                        procesar_parrafo(p)
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    procesar_parrafo(p)

        from lxml import etree
        for node in doc._element.xpath('.//w:t'):
            if not node.text: continue
        nombre_salida = f"{sigedo_corto}_oficio.docx"
        ruta_salida = get_base_dir() / "Informes_Generados" / nombre_salida
        ruta_salida.parent.mkdir(parents=True, exist_ok=True)
        
        log(f"Guardando Oficio en: {ruta_salida.name}")
        doc.save(ruta_salida)
        
        # Post-procesamiento: Reemplazo robusto en XML para texto dividido (como cuadros de texto)
        import zipfile
        import tempfile
        import os
        import re

        temp_dir = tempfile.mkdtemp()
        with zipfile.ZipFile(ruta_salida, 'r') as zin:
            zin.extractall(temp_dir)
            
        doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        if os.path.exists(doc_xml_path):
            with open(doc_xml_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Buscar 5...7...8...9...7
            patron_57897_2026 = re.compile(r'5(<[^>]+>)*7(<[^>]+>)*8(<[^>]+>)*9(<[^>]+>)*7(<[^>]+>)*-(<[^>]+>)*2(<[^>]+>)*0(<[^>]+>)*2(<[^>]+>)*6')
            content = patron_57897_2026.sub(reemplazo_sigedo, content)
            
            patron_57897 = re.compile(r'5(<[^>]+>)*7(<[^>]+>)*8(<[^>]+>)*9(<[^>]+>)*7')
            content = patron_57897.sub(sigedo_corto, content)
            
            with open(doc_xml_path, 'w', encoding='utf-8') as f:
                f.write(content)
                
        # Empaquetar de nuevo
        with zipfile.ZipFile(ruta_salida, 'w', zipfile.ZIP_DEFLATED) as zout:
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zout.write(file_path, arcname)
                    
        import shutil
        shutil.rmtree(temp_dir)
        
        return ruta_salida
