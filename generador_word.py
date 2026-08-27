from __future__ import annotations

import sys
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Callable

import docx
from docx import Document
from docxtpl import DocxTemplate


LogCallback = Callable[[str], None]


def get_base_dir() -> Path:
    """Devuelve el directorio base del script o ejecutable congelado con PyInstaller."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).parent
    return Path(__file__).resolve().parent


BASE_DIR = get_base_dir()


def obtener_ruta_plantilla() -> Path:
    """Busca y retorna la ruta existente de plantilla_informe.docx probando múltiples ubicaciones."""
    rutas_candidatas = [
        BASE_DIR / "plantillas" / "plantilla_informe.docx",
        BASE_DIR / "dist" / "main" / "plantillas" / "plantilla_informe.docx",
        BASE_DIR / "dist" / "plantillas" / "plantilla_informe.docx",
        Path(__file__).resolve().parent / "plantillas" / "plantilla_informe.docx",
    ]
    for ruta in rutas_candidatas:
        if ruta.is_file():
            return ruta
    return BASE_DIR / "plantillas" / "plantilla_informe.docx"


PLANTILLA_PATH = obtener_ruta_plantilla()
SALIDA_DIR = BASE_DIR / "Informes_Generados"


class GeneradorWord:
    """Generador híbrido ultra-robusto que aplica docxtpl Y reemplazo directo en python-docx
    cubriendo todas las variantes de sintaxis: {{VAR}}, [VAR], <VAR>, <<VAR>>, ${VAR}, {VAR}."""

    def generar(self, contexto: dict, log: LogCallback | None = None) -> Path:
        _log = log or (lambda msg: None)

        ruta_plantilla = obtener_ruta_plantilla()
        if not ruta_plantilla.exists():
            raise FileNotFoundError(
                f"No se encontró la plantilla en: {ruta_plantilla.resolve()}\n"
                "Coloca 'plantilla_informe.docx' dentro de la carpeta 'plantillas'."
            )

        SALIDA_DIR.mkdir(parents=True, exist_ok=True)

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        sigedo_val = str(contexto.get("NUMERO_SIGEDO", ""))
        sigedo_corto = sigedo_val.split("-")[0] if sigedo_val else f"informe_{ts}"
        nombre_archivo = f"{sigedo_corto}_informe.docx"
        ruta_salida = SALIDA_DIR / nombre_archivo

        _log(f"  Plantilla: {ruta_plantilla.resolve()}")
        _log(f"  Salida: {ruta_salida.resolve()}")

        # 1. Mapeo exhaustivo de claves y alias
        alias_map = {
            "NOMBRES_Y_APELLIDOS_VALIDADOS": [
                "NOMBRES_Y_APELLIDOS", "NOMBRE_Y_APELLIDOS", "APELLIDOS_Y_NOMBRES",
                "NOMBRE_BECARIO", "NOMBRE_BECARIA", "NOMBRE_ESTUDIANTE", "BECARIO",
                "BECARIA", "ESTUDIANTE", "ALUMNO", "ALUMNA", "NOMBRES", "APELLIDOS",
                "NOMBRE_COMPLETO", "DATOS_BECARIO", "BENEFICIARIO", "POSTULANTE"
            ],
            "DNI_VALIDADO": [
                "DNI", "DNI_BECARIO", "NUMERO_DNI", "NRO_DNI", "DOCUMENTO",
                "DOC_IDENTIDAD", "NUM_DNI", "NRO_DOC", "NUM_DOC"
            ],
            "BECA_Y_CONVOCATORIA_VALIDADA": [
                "BECA_Y_CONVOCATORIA", "PROGRAMA_Y_CONVOCATORIA", "BECA",
                "PROGRAMA_BECA", "PROGRAMA", "CONVOCATORIA", "MODALIDAD"
            ],
            "FECHA_INICIO_SIBEC": [
                "INICIO_SIBEC", "FECHA_INICIO", "F_INICIO", "INICIO_BECA"
            ],
            "FECHA_FIN_SIBEC": [
                "FIN_SIBEC", "FECHA_FIN", "F_FIN", "FIN_BECA"
            ],
            "FECHA_SOLICITUD_TEXTO": [
                "FECHA_SOLICITUD", "FECHA_SOL", "FECHA_INGRESO", "FECHA_REGISTRO"
            ],
            "NOMBRE_INFORME_SUCCOR": [
                "INFORME_SUCCOR", "INFORME", "NUMERO_INFORME", "NUM_INFORME", "NRO_INFORME"
            ],
            "NUMERO_SIGEDO": [
                "NUM_SIGEDO", "SIGEDO", "EXPEDIENTE_SIGEDO", "EXPEDIENTE"
            ],
            "SEMESTRE_SOLICITADO": [
                "SEMESTRE", "SEMESTRE_SOLICITUD", "PERIODO_SOLICITADO"
            ],
            "SEMESTRE_ANTERIOR": [
                "SEMESTRE_PREVIO", "PERIODO_ANTERIOR"
            ],
            "RJD_ADJUDICACION": [
                "RJD", "RESOLUCION", "RESOLUCION_RJD"
            ],
            "INSTITUCION": [
                "IES", "UNIVERSIDAD", "INSTITUTO", "INSTITUCION_EDUCATIVA"
            ],
            "CARRERA": [
                "ESPECIALIDAD", "PROGRAMA_ESTUDIOS", "CARRERA_PROFESIONAL"
            ],
            "FECHA_MATRICULA": [
                "MATRICULA", "FECHA_DE_MATRICULA"
            ],
            "FECHA_INICIO_ESTUDIOS": [
                "INICIO_ESTUDIOS", "FECHA_INICIO_CLASES"
            ],
            "ESTADO_PROCEDENCIA": [
                "PROCEDENCIA", "EVALUACION_PROCEDENCIA", "PROCEDENTE_OBSERVADO"
            ],
            "CODIGO_DOC_IES": [
                "CODIGO_IES", "CARTA_IES", "OFICIO_IES", "DOCUMENTO_IES"
            ],
            "FECHA_DOC_IES_TEXTO": [
                "FECHA_DOC_IES", "FECHA_CARTA_IES"
            ],
            "CURSOS_PENDIENTES": [
                "CURSOS", "CURSOS_PENDIENTES_TEXTO"
            ],
            "FECHA_ACTUAL_TEXTO": [
                "FECHA_ACTUAL", "FECHA_HOY", "HOY", "FECHA"
            ],
            "TABLA_CICLOS": [
                "CICLOS", "tabla_ciclos", "ciclos", "TABLA_DE_CICLOS"
            ],
            "EL_LA_BECARIO_A": [
                "SEXO_ARTICULO_2", "SEXO_BECARIO_A"
            ],
            "DEL_BECARIO_A": [
                "DEL_BECARIO", "DE_LA_BECARIA"
            ],
            "A_LA_BECARIO_A": [
                "AL_BECARIO", "A_LA_BECARIA"
            ],
            "BECA_TITULO": [
                "BECA_TITLE_CASE", "BECA_NOMBRE"
            ],
            "REFERENCIA_A": [
                "REF_A", "REFERENCIA_SOLICITUD"
            ],
            "REFERENCIA_B": [
                "REF_B", "REFERENCIA_INFORME"
            ],
            "NOMBRE_PRIMERO_NOMBRES": [
                "NOMBRE_TITULO", "NOMBRE_NOMBRES_PRIMERO"
            ],
        }

        # 2. Expandir contexto con claves y alias
        ctx_expandido = dict(contexto)
        for clave_principal, lista_alias in alias_map.items():
            if clave_principal in contexto:
                val = contexto[clave_principal]
                for alias in lista_alias:
                    if alias not in ctx_expandido:
                        ctx_expandido[alias] = val

        # 3. Preparar diccionario de reemplazos plano para cadenas (ignorando listas/dicts)
        # NOTA: CURSOS_PENDIENTES y sus alias se excluyen del mapa de texto plano porque
        # contienen saltos de línea y deben expandirse en párrafos reales por _expandir_cursos_pendientes (PASO 3).
        CLAVES_EXCLUIDAS_REEMPLAZO = {
            "CURSOS_PENDIENTES", "CURSOS", "CURSOS_PENDIENTES_TEXTO",
            "cursos_pendientes", "cursos", "cursos_pendientes_texto",
            "Cursos_Pendientes", "Cursos", "Cursos_Pendientes_Texto",
        }

        mapa_reemplazos = {}
        ctx_limpio = {}

        for k, v in ctx_expandido.items():
            if v is None:
                val_str = ""
                ctx_limpio[k] = ""
            elif isinstance(v, list):
                val_list = []
                for item in v:
                    if isinstance(item, dict):
                        item_var = dict(item)
                        for ik, iv in list(item.items()):
                            item_var[ik.lower()] = iv
                            item_var[ik.upper()] = iv
                            item_var[ik.title()] = iv
                        val_list.append(item_var)
                    else:
                        val_list.append(item)
                ctx_limpio[k] = val_list
                continue
            elif isinstance(v, dict):
                ctx_limpio[k] = v
                continue
            else:
                val_str = str(v)
                ctx_limpio[k] = val_str

            ctx_limpio[k.lower()] = val_str
            ctx_limpio[k.upper()] = val_str
            ctx_limpio[k.title()] = val_str

            # Excluir claves de cursos del mapa de reemplazo de texto plano
            if k.upper() not in CLAVES_EXCLUIDAS_REEMPLAZO and k not in CLAVES_EXCLUIDAS_REEMPLAZO:
                mapa_reemplazos[k.upper()] = val_str

        # PASO 1: Renderizado mediante docxtpl (Jinja2)
        # CURSOS_PENDIENTES se sustituye por un placeholder unico detectable.
        # Asi Jinja2 no lo borra ni lo convierte en texto plano con \n embebidos.
        # El PASO 3 (_expandir_cursos_pendientes) se encarga de expandirlo en parrafos reales.
        _PLACEHOLDER_CURSOS = "__CURSOS_PH__"
        ctx_tpl = {k: v for k, v in ctx_limpio.items() if k not in CLAVES_EXCLUIDAS_REEMPLAZO}
        ctx_tpl["CURSOS_PENDIENTES"] = _PLACEHOLDER_CURSOS
        ctx_tpl["cursos_pendientes"] = _PLACEHOLDER_CURSOS
        ctx_tpl["Cursos_Pendientes"] = _PLACEHOLDER_CURSOS
        try:
            tpl = DocxTemplate(ruta_plantilla)
            tpl.render(ctx_tpl)
            tpl.save(ruta_salida)
        except Exception as err:
            _log(f"  [AVISO docxtpl]: {err}. Se aplicará reemplazo directo con python-docx.")
            shutil.copy(ruta_plantilla, ruta_salida)

        # PASO 2: Reemplazo profundo directo con python-docx para TODAS las sintaxis y fallback de cadenas
        self._reemplazo_profundo_docx(ruta_salida, mapa_reemplazos, contexto)

        # PASO 3: Expansión especial de CURSOS_PENDIENTES en párrafos numerados individuales
        cursos_lista = contexto.get("CURSOS_PENDIENTES", "")
        if not cursos_lista:
            cursos_lista = "(No se detectaron cursos pendientes)"
        self._expandir_cursos_pendientes(ruta_salida, cursos_lista)

        _log("  Documento Word guardado correctamente.")
        return ruta_salida

    def _reemplazo_profundo_docx(self, ruta_docx: Path, mapa_reemplazos: dict[str, str], contexto: dict) -> None:
        """Abre el archivo Word y reemplaza cualquier etiqueta residual o cadena estática en párrafos, tablas, encabezados y pies."""
        doc = Document(ruta_docx)
        tabla_ciclos = contexto.get("TABLA_CICLOS", [])

        patrones_reemplazo = []
        for clave, valor in mapa_reemplazos.items():
            if not clave:
                continue
            nombres = {clave, clave.lower(), clave.title()}
            for n in nombres:
                patrones_reemplazo.extend([
                    (f"{{{{{n}}}}}", valor),
                    (f"{{{{ {n} }}}}", valor),
                    (f"[{n}]", valor),
                    (f"[ {n} ]", valor),
                    (f"<{n}>", valor),
                    (f"<<{n}>>", valor),
                    (f"${{{n}}}", valor),
                    (f"{{{n}}}", valor),
                ])

        # Mapeo de reemplazos estáticos de resguardo (para plantillas antiguas sin etiquetas)
        reemplazos_fallback = []
        if contexto.get("NOMBRES_Y_APELLIDOS_VALIDADOS"):
            val_nom = str(contexto["NOMBRES_Y_APELLIDOS_VALIDADOS"])
            reemplazos_fallback.extend([
                ("LAIDY SCARLE PANTOJA CUSI", val_nom),
                ("Laidy Scarle Pantoja Cusi", val_nom),
            ])
        if contexto.get("DNI_VALIDADO"):
            val_dni = str(contexto["DNI_VALIDADO"])
            reemplazos_fallback.append(("75551078", val_dni))
        if contexto.get("INSTITUCION"):
            val_ies = str(contexto["INSTITUCION"])
            reemplazos_fallback.append(("Universidad Peruana Cayetano Heredia", val_ies))
        if contexto.get("CODIGO_DOC_IES"):
            val_doc_ies = str(contexto["CODIGO_DOC_IES"])
            reemplazos_fallback.append(("CAR.OUB-UPCH-1565-2026", val_doc_ies))
        if contexto.get("NOMBRE_INFORME_SUCCOR"):
            val_succor = str(contexto["NOMBRE_INFORME_SUCCOR"])
            reemplazos_fallback.extend([
                ("Informe Nº 4187-2026-MINEDU/VMGI-PRONABEC-DICONCI-SUCCOR-LIMA", val_succor),
                ("Informe N° 4187-2026-MINEDU/VMGI-PRONABEC-DICONCI-SUCCOR-LIMA", val_succor),
                ("INFORME N° 4187-2026-MINEDU/VMGI-PRONABEC-DICONCI-SUCCOR-LIMA", val_succor),
            ])
        if contexto.get("CARRERA"):
            val_carr = str(contexto["CARRERA"])
            reemplazos_fallback.append(("CARRERA PROFESIONAL DE ENFERMERIA", val_carr))
        if contexto.get("BECA_Y_CONVOCATORIA_VALIDADA"):
            val_beca = str(contexto["BECA_Y_CONVOCATORIA_VALIDADA"])
            reemplazos_fallback.append(("BECA 18 - 2021", val_beca))
            
            val_sexo = str(contexto.get("SEXO_ARTICULO_1", "1 becario/a"))
            val_beca_titulo = str(contexto.get("BECA_TITULO", val_beca))
            
            reemplazos_fallback.extend([
                ("1 becaria de la Beca 18 - Convocatoria 2021", f"{val_sexo} de la {val_beca_titulo}"),
                ("1 becario de la Beca 18 - Convocatoria 2021", f"{val_sexo} de la {val_beca_titulo}"),
            ])
        if contexto.get("FECHA_SOLICITUD_TEXTO"):
            val_fsol = str(contexto["FECHA_SOLICITUD_TEXTO"])
            reemplazos_fallback.extend([
                ("17 DE JULIO DE 2026", val_fsol),
                ("17 de julio de 2026", val_fsol),
            ])
        if contexto.get("FECHA_DOC_IES_TEXTO"):
            val_fdoc = str(contexto["FECHA_DOC_IES_TEXTO"])
            reemplazos_fallback.append(("16 de julio del 2026", val_fdoc))
        if contexto.get("FECHA_INICIO_SIBEC"):
            val_fini = str(contexto["FECHA_INICIO_SIBEC"])
            reemplazos_fallback.extend([
                ("23-08-2021", val_fini),
                ("23/08/2021", val_fini),
            ])
        if contexto.get("FECHA_FIN_SIBEC"):
            val_ffin = str(contexto["FECHA_FIN_SIBEC"])
            reemplazos_fallback.extend([
                ("14-09-2026", val_ffin),
                ("14/09/2026", val_ffin),
            ])

        def _reemplazar_texto_run(run, patrones_reemplazo, reemplazos_fallback, nro_inf_gen):
            """Reemplaza texto dentro de un run preservando el formato original."""
            texto_nuevo = run.text
            for tag, val in patrones_reemplazo:
                if tag in texto_nuevo:
                    texto_nuevo = texto_nuevo.replace(tag, val)
            for old_val, new_val in reemplazos_fallback:
                if old_val in texto_nuevo:
                    texto_nuevo = texto_nuevo.replace(old_val, new_val)
            if nro_inf_gen:
                texto_nuevo = re.sub(r"(INFORME\s+N[º°o]?\s*)\d+(-\d{4}-MINEDU/VMGI-PRONABEC-DIBEC-SUS)", rf"\g<1>{nro_inf_gen}\g<2>", texto_nuevo, flags=re.IGNORECASE)
            if texto_nuevo != run.text:
                run.text = texto_nuevo

        def procesar_parrafo(p):
            if not p.text:
                return
            nro_inf_gen = contexto.get("NUMERO_INFORME_GENERAR", "")
            
            # Primero intentar reemplazo a nivel de run (preserva formato de cada run)
            texto_pre = p.text
            for run in p.runs:
                _reemplazar_texto_run(run, patrones_reemplazo, reemplazos_fallback, nro_inf_gen)

            # Fallback: si el párrafo sigue teniendo etiquetas residuales
            # (etiqueta dividida entre varios runs), aplicar a nivel de párrafo
            if p.text == texto_pre:
                texto_orig = p.text
                texto_nuevo = texto_orig
                for tag, val in patrones_reemplazo:
                    if tag in texto_nuevo:
                        texto_nuevo = texto_nuevo.replace(tag, val)
                for old_val, new_val in reemplazos_fallback:
                    if old_val in texto_nuevo:
                        texto_nuevo = texto_nuevo.replace(old_val, new_val)
                if nro_inf_gen:
                    texto_nuevo = re.sub(r"(INFORME\s+N[º°o]?\s*)\d+(-\d{4}-MINEDU/VMGI-PRONABEC-DIBEC-SUS)", rf"\g<1>{nro_inf_gen}\g<2>", texto_nuevo, flags=re.IGNORECASE)
                
                if texto_nuevo != texto_orig:
                    # Guardar formato del primer run antes de sobrescribir
                    fmt_guardado = {}
                    if p.runs:
                        r0 = p.runs[0]
                        fmt_guardado = {
                            "name": r0.font.name,
                            "size": r0.font.size,
                            "bold": r0.bold,
                            "italic": r0.italic,
                            "underline": r0.underline,
                        }
                    p.text = texto_nuevo
                    # Restaurar formato al nuevo run para mantener Arial 11
                    if p.runs and fmt_guardado:
                        r = p.runs[0]
                        if fmt_guardado.get("name"):
                            r.font.name = fmt_guardado["name"]
                        if fmt_guardado.get("size"):
                            r.font.size = fmt_guardado["size"]
                        if fmt_guardado.get("bold") is not None:
                            r.bold = fmt_guardado["bold"]
                        if fmt_guardado.get("italic") is not None:
                            r.italic = fmt_guardado["italic"]
                        if fmt_guardado.get("underline") is not None:
                            r.underline = fmt_guardado["underline"]

        # 1. Párrafos principales
        for p in doc.paragraphs:
            procesar_parrafo(p)

        # 2. Tablas principales
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        procesar_parrafo(p)

        # 3. Encabezados y pies de página
        for section in doc.sections:
            for p in section.header.paragraphs:
                procesar_parrafo(p)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            procesar_parrafo(p)

            for p in section.footer.paragraphs:
                procesar_parrafo(p)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            procesar_parrafo(p)

        # 4. Rellenado dinámico de Tabla de Ciclos
        # Detectar la tabla por columnas Momento/Ciclo/Semestre y reconstruirla
        # con 'Duración de estudios' fusionado en columna 0 y ordinales en columna 1.
        if tabla_ciclos and isinstance(tabla_ciclos, list) and len(tabla_ciclos) > 0:
            from docx.shared import Pt as _Pt2
            from docx.enum.text import WD_ALIGN_PARAGRAPH as _WAP
            from docx.enum.table import WD_ALIGN_VERTICAL as _WAV

            for table in doc.tables:
                if len(table.rows) == 0:
                    continue
                first_row_texts = [cell.text.strip().upper() for cell in table.rows[0].cells]
                tiene_momento = any("MOMENTO" in t for t in first_row_texts)
                tiene_ciclo = any("CICLO" in t for t in first_row_texts)
                tiene_semestre = any("SEMESTRE" in t for t in first_row_texts)
                if not (tiene_momento and tiene_ciclo and tiene_semestre):
                    continue

                # Limpiar filas de datos (mantener solo encabezado)
                while len(table.rows) > 1:
                    tr = table.rows[-1]._tr
                    table._tbl.remove(tr)

                # Agregar filas de ciclos (col 0 vacía, col 1 = ordinal, col 2 = semestre)
                for item in tabla_ciclos:
                    new_row = table.add_row()
                    nc0 = new_row.cells[0] if len(new_row.cells) > 0 else None
                    nc1 = new_row.cells[1] if len(new_row.cells) > 1 else None
                    nc2 = new_row.cells[2] if len(new_row.cells) > 2 else None

                    val_ciclo = str(
                        item.get("ciclo") or item.get("CICLO") or item.get("Ciclo") or ""
                    )
                    val_semestre = str(
                        item.get("semestre") or item.get("SEMESTRE") or item.get("Semestre") or ""
                    )

                    if nc0:
                        nc0.text = ""
                        nc0.vertical_alignment = _WAV.CENTER
                    if nc1:
                        nc1.text = val_ciclo
                        nc1.vertical_alignment = _WAV.CENTER
                        for para in nc1.paragraphs:
                            para.alignment = _WAP.CENTER
                            for run in para.runs:
                                run.font.name = "Arial"
                                run.font.size = _Pt2(11)
                    if nc2:
                        nc2.text = val_semestre
                        nc2.vertical_alignment = _WAV.CENTER
                        for para in nc2.paragraphs:
                            para.alignment = _WAP.CENTER
                            for run in para.runs:
                                run.font.name = "Arial"
                                run.font.size = _Pt2(11)

                # Fusionar celdas de columna 0 en todas las filas de datos
                n_datos = len(tabla_ciclos)
                if n_datos >= 1:
                    from docx.oxml.ns import qn as _qn2
                    celda_inicio = table.cell(1, 0)
                    if n_datos > 1:
                        celda_fin = table.cell(n_datos, 0)
                        celda_inicio.merge(celda_fin)
                    # Escribir 'Duración de estudios' en la celda fusionada
                    celda_inicio.text = "Duración de estudios"
                    celda_inicio.vertical_alignment = _WAV.CENTER
                    for para in celda_inicio.paragraphs:
                        para.alignment = _WAP.CENTER
                        for run in para.runs:
                            run.font.name = "Arial"
                            run.font.size = _Pt2(11)
                    if not celda_inicio.paragraphs[0].runs:
                        run = celda_inicio.paragraphs[0].add_run("Duración de estudios")
                        run.font.name = "Arial"
                        run.font.size = _Pt2(11)

                break  # Solo procesar la primera tabla de ciclos encontrada

        doc.save(ruta_docx)

    def _expandir_cursos_pendientes(self, ruta_docx: Path, cursos_texto: str) -> None:
        """Detecta en el documento Word el párrafo que contiene el listado de cursos
        pendientes (marcador o texto ya renderizado con saltos de línea) y lo expande
        en párrafos individuales numerados, clonando el formato XML del párrafo original."""
        import copy
        from lxml import etree
        from docx.oxml.ns import qn

        doc = Document(ruta_docx)

        # Construir lista de líneas de cursos (quitar vacíos)
        lineas = [l.strip() for l in cursos_texto.splitlines() if l.strip()]
        if not lineas:
            doc.save(ruta_docx)
            return

        # Asegurarse de que cada línea tenga su numeración
        lineas_numeradas = []
        for i, linea in enumerate(lineas):
            # Si la línea ya empieza con "N." (ej. "1. Curso"), usarla tal cual
            if re.match(r"^\d+\.\s+", linea):
                lineas_numeradas.append(linea)
            else:
                lineas_numeradas.append(f"{i+1}. {linea}")

        # Buscar el párrafo en el cuerpo que contiene el marcador residual o
        # que ya tiene el texto completo de cursos (todos juntos en un solo párrafo)
        body = doc.element.body
        parrafos_body = [p for p in body.iterchildren() if p.tag.endswith('}p')]

        idx_objetivo = None
        parrafo_objetivo = None

        MARCADORES = (
            "__CURSOS_PH__",          # placeholder inyectado por PASO 1
            "{{ CURSOS_PENDIENTES }}",
            "{{CURSOS_PENDIENTES}}",
            "[ CURSOS_PENDIENTES ]",
            "[CURSOS_PENDIENTES]",
            "<CURSOS_PENDIENTES>",
            "${CURSOS_PENDIENTES}",
            "{CURSOS_PENDIENTES}",
            "{{ cursos_pendientes }}",
            "{{ Cursos_Pendientes }}",
        )

        for idx, p_elem in enumerate(parrafos_body):
            # Obtener texto completo del párrafo
            texto_p = "".join(
                t.text or "" for t in p_elem.iter(qn("w:t"))
            ).strip()
            # Detectar marcador residual O placeholder O texto ya renderizado con el primer curso
            es_marcador = any(m in texto_p for m in MARCADORES)
            # Detectar texto ya renderizado: contiene la primera línea del listado
            primera_linea = lineas_numeradas[0] if lineas_numeradas else ""
            es_renderizado = primera_linea and primera_linea in texto_p and len(texto_p) > len(primera_linea)
            es_solo_primera = primera_linea and texto_p == primera_linea and len(lineas_numeradas) > 1

            if es_marcador or es_renderizado or es_solo_primera:
                idx_objetivo = idx
                parrafo_objetivo = p_elem
                break

        # Eliminar parrafos vacios con numId (lista) que sigan al parrafo objetivo.
        # La plantilla tiene un P_extra vacio [numId=32] despues de {{ CURSOS_PENDIENTES }}
        # que generaria un numero extra en blanco.
        if parrafo_objetivo is not None:
            siguiente = parrafo_objetivo.getnext()
            while siguiente is not None and siguiente.tag.endswith('}p'):
                # Solo eliminar si es un parrafo de lista numerada (tiene numPr) y esta vacio
                p_texto = "".join(t.text or "" for t in siguiente.iter(qn("w:t"))).strip()
                p_pPr = siguiente.find(qn("w:pPr"))
                p_numPr = p_pPr.find(qn("w:numPr")) if p_pPr is not None else None
                if p_numPr is not None and not p_texto:
                    siguiente_sig = siguiente.getnext()
                    siguiente.getparent().remove(siguiente)
                    siguiente = siguiente_sig
                else:
                    break

        if parrafo_objetivo is None or idx_objetivo is None:
            # No se encontró el párrafo objetivo; no modificar
            doc.save(ruta_docx)
            return

        # Clonar el párrafo plantilla (formato: estilo, numeración, indentación, fuente)
        parrafo_base = copy.deepcopy(parrafo_objetivo)

        # Limpiar todos los runs del párrafo base clonado para usarlo como molde
        for r_elem in parrafo_base.findall(qn("w:r")):
            parrafo_base.remove(r_elem)
        # También limpiar w:hyperlink y otros hijos que no sean w:pPr
        for child in list(parrafo_base):
            if child.tag not in (qn("w:pPr"),):
                parrafo_base.remove(child)

        # Obtener el rPr (formato de caracteres) del run original si existe
        rPr_original = None
        runs_orig = parrafo_objetivo.findall(qn("w:r"))
        if runs_orig:
            rPr_original = runs_orig[0].find(qn("w:rPr"))

        def crear_parrafo_curso(texto_curso: str) -> etree._Element:
            """Crea un elemento <w:p> clonado del párrafo base con el texto del curso."""
            p_nuevo = copy.deepcopy(parrafo_base)
            # Crear el run con el texto
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            r_elem = etree.SubElement(p_nuevo, qn("w:r"))
            if rPr_original is not None:
                r_elem.append(copy.deepcopy(rPr_original))
            t_elem = etree.SubElement(r_elem, qn("w:t"))
            t_elem.text = texto_curso
            t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            return p_nuevo

        # Reemplazar el párrafo objetivo por los N párrafos de cursos
        # Primero: reemplazar el contenido del párrafo objetivo con el primer curso
        # Luego: insertar los demás inmediatamente después

        # Limpiar runs del párrafo objetivo
        for r_elem in list(parrafo_objetivo.findall(qn("w:r"))):
            parrafo_objetivo.remove(r_elem)
        for child in list(parrafo_objetivo):
            if child.tag not in (qn("w:pPr"),):
                parrafo_objetivo.remove(child)

        # Insertar texto del primer curso en el párrafo objetivo
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        r0 = etree.SubElement(parrafo_objetivo, qn("w:r"))
        if rPr_original is not None:
            r0.append(copy.deepcopy(rPr_original))
        t0 = etree.SubElement(r0, qn("w:t"))
        t0.text = lineas_numeradas[0]
        t0.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

        # Insertar los párrafos restantes después del párrafo objetivo
        insert_after = parrafo_objetivo
        for linea in lineas_numeradas[1:]:
            p_nuevo = crear_parrafo_curso(linea)
            insert_after.addnext(p_nuevo)
            insert_after = p_nuevo

        doc.save(ruta_docx)

        # 5. Aplicar formato Arial 8 + MAYÚsCULAS + centrado vertical a la tabla del becario
        # (la que contiene columnas como DNI, BECARIO/A, RJ, BECA Y CONVOCATORIA, INSTITUCIÓN, CARRERA)
        from docx.oxml.ns import qn as _qn
        from docx.shared import Pt as _Pt
        from docx.enum.table import WD_ALIGN_VERTICAL

        doc2 = Document(ruta_docx)
        PALABRAS_TABLA_BECARIO = ("DNI", "BECARIO", "BECA Y CONVOCATORIA", "INSTITUCION", "CARRERA", "CONVOCATORIA", "RJD", "RJ ")
        for table in doc2.tables:
            encabezado = " ".join(cell.text for row in table.rows[:2] for cell in row.cells).upper()
            if sum(1 for w in PALABRAS_TABLA_BECARIO if w in encabezado) >= 2:
                for row in table.rows:
                    for cell in row.cells:
                        # Centrado vertical
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.name = "Arial"
                                run.font.size = _Pt(8)
                                run.text = run.text.upper()
                            # Si no hay runs pero hay texto en el párrafo
                            if not para.runs and para.text.strip():
                                run = para.add_run(para.text.upper())
                                run.font.name = "Arial"
                                run.font.size = _Pt(8)
                                # Limpiar el texto del párrafo original
                                for child in list(para._p):
                                    if child.tag != _qn("w:r"):
                                        continue
                                    if child is not run._r:
                                        para._p.remove(child)
                break
        doc2.save(ruta_docx)

